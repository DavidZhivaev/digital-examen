import hashlib
import math
import re
import uuid

from fastapi import HTTPException
from tortoise.transactions import in_transaction

from core.config import settings
from tasks.models import Task, TaskBank, TaskPosition, TaskReview, TaskRevision, generate_task_code
from tasks.permissions import can_manage_bank, teacher_teaches_parallel


TASK_STATUS_DRAFT = 0
TASK_STATUS_PENDING = 1
TASK_STATUS_APPROVED = 2
TASK_STATUS_REJECTED = 3


def validate_latex_content(*values: str | None):
    for value in values:
        if not value:
            continue

        if value.count("$") % 2 != 0:
            raise HTTPException(400, "LaTeX содержит незакрытый символ $")

        pairs = [("{", "}"), ("[", "]"), ("(", ")")]
        for left, right in pairs:
            if value.count(left) != value.count(right):
                raise HTTPException(400, f"LaTeX содержит несбалансированные скобки {left}{right}")

        for command in re.findall(r"\\[A-Za-z]+", value):
            if len(command) == 1:
                raise HTTPException(400, "LaTeX содержит битую команду")


def stable_hash(bank_id: int, task_id: uuid.UUID) -> int:
    key = f"{bank_id}:{str(task_id)}"
    return int(hashlib.md5(key.encode()).hexdigest(), 16)


def is_public_task(bank: TaskBank, task: Task) -> bool:
    if not bank.is_global or not bank.is_open:
        return False
    if task.status != TASK_STATUS_APPROVED:
        return False
    return stable_hash(bank.id, task.id) % 100 < bank.visibility_percent


def paginate(items: list, page: int = 1, limit: int = 10) -> dict:
    page = max(page, 1)
    limit = max(1, min(limit, 100))
    total = len(items)
    start = (page - 1) * limit
    end = start + limit

    return {
        "page": page,
        "limit": limit,
        "total": total,
        "pages": math.ceil(total / limit) if total else 0,
        "items": items[start:end],
    }


async def generate_unique_task_code() -> str:
    for _ in range(20):
        code = generate_task_code()
        if not await Task.filter(code=code).exists():
            return code

    raise HTTPException(500, "Не удалось сгенерировать уникальный номер задачи")


async def create_revision(task: Task, user):
    defaults = {
        "text": task.text,
        "solution": task.solution,
        "answer": task.answer,
        "image_url": task.image_url,
        "image_scale": task.image_scale,
        "image_position": task.image_position,
        "status": task.status,
        "changed_by": user,
    }

    await TaskRevision.update_or_create(
        defaults=defaults,
        task=task,
        version=task.version,
    )


async def create_review(task: Task, moderator, action: str, comment: str | None = None):
    await TaskReview.create(
        task=task,
        moderator=moderator,
        action=action,
        comment=comment,
    )


async def latest_review_comments(task_ids: list[uuid.UUID]) -> dict[str, str | None]:
    if not task_ids:
        return {}

    reviews = await TaskReview.filter(task_id__in=task_ids).order_by("task_id", "-created_at")
    result: dict[str, str | None] = {}

    for review in reviews:
        key = str(review.task_id)
        if key not in result:
            result[key] = review.comment

    return result


async def reorder_positions(bank: TaskBank, new_order: list[int]):
    async with in_transaction():
        positions = await bank.positions.all()
        pos_map = {p.id: p for p in positions}

        if len(new_order) != len(positions) or set(new_order) != set(pos_map.keys()):
            raise HTTPException(400, "Передан некорректный список ID для сортировки")

        for position in positions:
            position.order = position.order + 10000
            await position.save()

        for idx, position_id in enumerate(new_order, start=1):
            position = pos_map[position_id]
            position.order = idx
            await position.save()


def convert_docx_to_math_text(file_path: str) -> str:
    from docx import Document

    document = Document(file_path)
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


def convert_pdf_to_math_text(file_path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HTTPException(
            400,
            "Распознавание PDF требует установленный пакет pypdf",
        ) from exc

    reader = PdfReader(file_path)
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


class TaskBankService:
    @staticmethod
    async def init_positions(bank: TaskBank, positions_count: int):
        for order in range(1, positions_count + 1):
            await TaskPosition.create(bank=bank, order=order)

    @staticmethod
    async def sync_positions_count(bank: TaskBank, positions_count: int):
        current_positions = await TaskPosition.filter(bank=bank).order_by("order")
        current_count = len(current_positions)

        if positions_count > current_count:
            for order in range(current_count + 1, positions_count + 1):
                await TaskPosition.create(bank=bank, order=order)

        if positions_count < current_count:
            extra_positions = [p for p in current_positions if p.order > positions_count]
            for position in extra_positions:
                has_tasks = await Task.filter(position=position, is_deleted=False).exists()
                if has_tasks:
                    raise HTTPException(
                        400,
                        "Нельзя уменьшить количество позиций: в удаляемых позициях есть задачи",
                    )
                await position.delete()

        bank.positions_count = positions_count
        await bank.save()

    @staticmethod
    async def available_tasks_count(bank: TaskBank, user=None) -> int:
        tasks = await Task.filter(
            position__bank_id=bank.id,
            is_deleted=False,
        ).prefetch_related("position")

        if user is None:
            return len([task for task in tasks if is_public_task(bank, task)])

        context = await TaskVisibilityService.get_bank_context(user, bank)
        return len(TaskVisibilityService.filter_visible_tasks(tasks, user, bank, context))

    @staticmethod
    async def full_response(bank: TaskBank, user=None) -> dict:
        subject_name = getattr(getattr(bank, "subject", None), "name", None)

        return {
            "id": bank.id,
            "title": bank.title,
            "subject_id": bank.subject_id,
            "subject_name": subject_name,
            "parallel": bank.parallel,
            "is_global": bank.is_global,
            "is_open": bank.is_open,
            "visibility_percent": bank.visibility_percent,
            "positions_count": bank.positions_count,
            "created_by_id": bank.created_by_id,
            "available_tasks_count": await TaskBankService.available_tasks_count(bank, user),
        }


class TaskVisibilityService:
    @staticmethod
    async def get_user_context(user, subject) -> tuple[bool, bool, bool]:
        is_operator = user.role >= settings.OPERATOR_ROLE
        is_admin = await subject.admins.filter(id=user.id).exists()
        is_teacher = await subject.teachers.filter(id=user.id).exists()
        return is_operator, is_admin, is_teacher

    @staticmethod
    async def get_bank_context(user, bank: TaskBank) -> dict:
        subject = bank.subject
        is_operator = user.role >= settings.OPERATOR_ROLE
        is_subject_admin = await subject.admins.filter(id=user.id).exists()
        is_subject_teacher = await subject.teachers.filter(id=user.id).exists()
        manages_bank = await can_manage_bank(user, bank)
        teaches_parallel = False

        if is_subject_teacher:
            teaches_parallel = await teacher_teaches_parallel(user, subject.id, bank.parallel)

        return {
            "is_operator": is_operator,
            "is_subject_admin": is_subject_admin,
            "is_subject_teacher": is_subject_teacher,
            "is_bank_admin": manages_bank,
            "teaches_parallel": teaches_parallel,
            "normal_view": manages_bank or teaches_parallel,
            "public_only": not manages_bank,
        }

    @staticmethod
    def filter_visible_tasks(tasks: list[Task], user, bank: TaskBank, context: dict) -> list[Task]:
        if context["is_bank_admin"]:
            return [task for task in tasks if not task.is_deleted]

        visible_tasks: list[Task] = []

        for task in tasks:
            if task.is_deleted:
                continue

            is_own_task = task.author_id == user.id
            if is_own_task:
                visible_tasks.append(task)
                continue

            if is_public_task(bank, task):
                visible_tasks.append(task)

        return visible_tasks

    @staticmethod
    async def serialize_tasks(tasks: list[Task], user, bank: TaskBank, context: dict) -> list[dict]:
        visible_tasks = TaskVisibilityService.filter_visible_tasks(tasks, user, bank, context)

        if context["normal_view"]:
            visible_tasks.sort(key=lambda task: (TaskVisibilityService._position_order(task), task.code))
        else:
            visible_tasks.sort(key=lambda task: (TaskVisibilityService._position_order(task), task.code))

        review_map = await latest_review_comments([task.id for task in visible_tasks])

        return [
            TaskVisibilityService._to_dict(
                task,
                show_answers=context["is_bank_admin"] or task.author_id == user.id,
                show_meta=context["normal_view"] or task.author_id == user.id,
                latest_review_comment=review_map.get(str(task.id)),
            )
            for task in visible_tasks
        ]

    @staticmethod
    def filter_and_serialize(
        tasks: list[Task],
        user,
        bank: TaskBank,
        is_op: bool,
        is_adm: bool,
        is_tech: bool,
    ) -> list[dict]:
        context = {
            "is_bank_admin": is_op or is_adm,
            "normal_view": is_op or is_adm or is_tech,
        }
        visible_tasks = TaskVisibilityService.filter_visible_tasks(tasks, user, bank, context)
        return [
            TaskVisibilityService._to_dict(
                task,
                show_answers=(is_op or is_adm or task.author_id == user.id),
                show_meta=(is_op or is_adm or is_tech or task.author_id == user.id),
            )
            for task in visible_tasks
        ]

    @staticmethod
    def _position_order(task: Task) -> int:
        position = getattr(task, "position", None)
        return getattr(position, "order", 0) or 0

    @staticmethod
    def _to_dict(
        task: Task,
        show_answers: bool,
        show_meta: bool = False,
        latest_review_comment: str | None = None,
    ) -> dict:
        data = {
            "id": str(task.id),
            "code": task.code,
            "position_id": task.position_id,
            "position_order": TaskVisibilityService._position_order(task),
            "text": task.text,
            "image_url": task.image_url,
            "image_scale": task.image_scale,
            "image_position": task.image_position,
        }

        if show_answers:
            data["solution"] = task.solution
            data["answer"] = task.answer
            data["author_id"] = task.author_id

        if show_meta:
            data["status"] = task.status
            data["version"] = task.version
            data["latest_review_comment"] = latest_review_comment

        return data
